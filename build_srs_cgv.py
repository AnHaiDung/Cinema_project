from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUTPUT = "SRS_CGV_Cinema.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
        if widths:
            hdr[i].width = widths[i]

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def add_kv_table(doc, rows):
    return add_table(doc, ["Thuộc tính", "Nội dung"], rows, [Cm(4), Cm(12.5)])


def add_mermaid(doc, title, code):
    doc.add_heading(title, level=3)
    p = doc.add_paragraph()
    p.add_run("Mã Mermaid dùng để dán vào draw.io:")
    for line in code.strip().splitlines():
        para = doc.add_paragraph()
        para.style = "Mermaid Code"
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_numbered_lines(doc, lines):
    for idx, line in enumerate(lines, start=1):
        doc.add_paragraph(f"{idx}. {line}", style="List Number")


def add_use_case(doc, number, uc_id, name, actor, trigger, pre, post, flow, fields):
    doc.add_heading(f"3.{number}. {name}", level=2)
    add_kv_table(
        doc,
        [
            ("Use Case ID", uc_id),
            ("Mô tả", name),
            ("Tác nhân (Actor)", actor),
            ("Sự ưu tiên (Priority)", "Cao"),
            ("Trigger", trigger),
            ("Điều kiện cần (Pre-Condition)", pre),
            ("Điều kiện sau (Post-Condition)", post),
        ],
    )
    doc.add_heading(f"3.{number}.1. Luồng cơ bản", level=3)
    add_numbered_lines(doc, flow)
    doc.add_heading(f"3.{number}.2. Giao diện", level=3)
    doc.add_paragraph(
        "Hệ thống sử dụng giao diện web Thymeleaf tương ứng với chức năng. "
        "Ảnh minh họa màn hình có thể bổ sung sau khi giao diện được chốt."
    )
    doc.add_heading(f"3.{number}.3. Mô tả dữ liệu nhập/xuất", level=3)
    add_table(doc, ["Tên kỹ thuật", "Tên tiếng Việt", "Loại", "Bắt buộc?", "Mô tả"], fields,
              [Cm(3.2), Cm(3.2), Cm(2.4), Cm(2.1), Cm(5.6)])


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    code = styles.add_style("Mermaid Code", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.left_indent = Cm(0.45)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.49)
        sec.footer_distance = Inches(0.49)

    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TÀI LIỆU ĐẶC TẢ NGHIỆP VỤ")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Hệ thống đặt vé và quản lý rạp phim CGV Cinema")
    r.bold = True
    r.font.size = Pt(14)

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project.add_run("DỰ ÁN: CGV Cinema Web Application").bold = True

    doc.add_paragraph()
    add_table(
        doc,
        ["Thông tin", "Nội dung"],
        [
            ("Người thực hiện", "Admin"),
            ("Ngày ban hành", "21/07/2026"),
            ("Mã tài liệu", "SRS-CGV-CINEMA-001"),
            ("Phiên bản", "1.0"),
            ("Công nghệ chính", "Spring Boot, Thymeleaf, Spring Security, Spring Data JPA, H2, Maven"),
        ],
        [Cm(5), Cm(11.5)],
    )

    doc.add_page_break()
    doc.add_heading("MỤC LỤC", level=1)
    toc_lines = [
        "PHẦN 1: GIỚI THIỆU",
        "1.1 Mục đích tài liệu",
        "1.2 Phạm vi tài liệu",
        "1.3 Tổng quan ứng dụng",
        "1.4 Thuật ngữ viết tắt",
        "PHẦN 2: YÊU CẦU TỔNG THỂ",
        "2.1 Sơ đồ quan hệ đối tượng",
        "2.2 Sơ đồ Use Case",
        "2.3 Sơ đồ luồng tổng thể",
        "2.4 Sơ đồ chuyển trạng thái",
        "2.5 Phân quyền",
        "2.6 Site Map",
        "PHẦN 3: CHỨC NĂNG",
        "PHẦN 4: COMPONENT, THÔNG BÁO, CẢNH BÁO",
        "PHẦN 5: LINK ISSUE",
    ]
    for line in toc_lines:
        doc.add_paragraph(line)

    doc.add_page_break()
    doc.add_heading("PHẦN 1: GIỚI THIỆU", level=1)
    doc.add_heading("1.1 Mục đích tài liệu", level=2)
    doc.add_paragraph(
        "Tài liệu này mô tả yêu cầu nghiệp vụ và yêu cầu chức năng cho hệ thống web đặt vé "
        "và quản lý rạp phim CGV Cinema. Tài liệu là cơ sở để phân tích, thiết kế, phát triển, "
        "kiểm thử và nghiệm thu sản phẩm."
    )
    doc.add_heading("1.2 Phạm vi tài liệu", level=2)
    scope = [
        "Khách hàng xem danh sách phim, xem chi tiết phim và các suất chiếu đang bán.",
        "Khách hàng đăng ký, đăng nhập, chọn ghế, đặt vé và xem lịch sử đặt vé.",
        "Quản trị viên quản lý danh mục phim, poster phim, phòng chiếu và lịch chiếu.",
        "Hệ thống phân quyền theo hai vai trò chính: ROLE_ADMIN và ROLE_CUSTOMER.",
        "Các sơ đồ nghiệp vụ được trình bày bằng mã Mermaid để có thể dán vào draw.io.",
        "Chưa bao gồm thanh toán trực tuyến, hoàn/hủy vé, sơ đồ ghế cấu hình động theo phòng, khuyến mãi và báo cáo doanh thu nâng cao.",
    ]
    for item in scope:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.3 Tổng quan ứng dụng", level=2)
    doc.add_paragraph(
        "Ứng dụng được xây dựng bằng Spring Boot theo mô hình MVC. Frontend sử dụng Thymeleaf template "
        "và CSS tĩnh. Backend gồm controller, service, repository và entity JPA. Dữ liệu demo sử dụng H2 "
        "in-memory; poster phim được lưu trong thư mục uploads/movies và được ánh xạ thành tài nguyên tĩnh."
    )

    doc.add_heading("1.4 Thuật ngữ viết tắt", level=2)
    add_table(
        doc,
        ["STT", "Từ viết tắt", "Diễn giải"],
        [
            ("1", "SRS", "Software Requirements Specification - tài liệu đặc tả yêu cầu phần mềm."),
            ("2", "CRUD", "Create, Read, Update, Delete - thêm, xem, sửa, xóa dữ liệu."),
            ("3", "MVC", "Model - View - Controller."),
            ("4", "JPA", "Java Persistence API, dùng để ánh xạ entity với cơ sở dữ liệu."),
            ("5", "UI", "User Interface - giao diện người dùng."),
            ("6", "RBAC", "Role-Based Access Control - kiểm soát truy cập theo vai trò."),
        ],
        [Cm(1.5), Cm(3.2), Cm(11.8)],
    )

    doc.add_heading("PHẦN 2: YÊU CẦU TỔNG THỂ", level=1)
    doc.add_heading("2.1 Sơ đồ quan hệ đối tượng", level=2)
    doc.add_paragraph("Sơ đồ mô tả các entity chính và quan hệ dữ liệu trong hệ thống.")
    add_mermaid(doc, "2.1.1 Mermaid - ERD", """
erDiagram
    USERS ||--o{ BOOKINGS : places
    MOVIES ||--o{ SHOWTIMES : has
    AUDITORIUMS ||--o{ SHOWTIMES : hosts
    SHOWTIMES ||--o{ BOOKINGS : receives

    USERS {
        Long id PK
        String username UK
        String password
        String fullName
        String email UK
        String role
    }
    MOVIES {
        Long id PK
        String title
        String genre
        Integer durationMinutes
        String ageRating
        LocalDate releaseDate
        String posterUrl
        String description
    }
    AUDITORIUMS {
        Long id PK
        String name UK
        Integer capacity
        String screenType
    }
    SHOWTIMES {
        Long id PK
        Long movie_id FK
        Long auditorium_id FK
        LocalDateTime startsAt
        Double ticketPrice
        ShowtimeStatus status
        String note
    }
    BOOKINGS {
        Long id PK
        Long user_id FK
        Long showtime_id FK
        String seats
        Double totalPrice
        LocalDateTime bookingDate
    }
""")

    doc.add_heading("2.2 Sơ đồ Use Case", level=2)
    doc.add_paragraph("Tác nhân chính gồm khách hàng và quản trị viên rạp phim.")
    add_mermaid(doc, "2.2.1 Mermaid - Use Case", """
flowchart LR
    Customer[Khách hàng]
    Admin[Quản trị viên]

    subgraph Public[Chức năng công khai]
        UC01[Xem danh sách phim]
        UC02[Xem chi tiết phim]
        UC03[Đăng ký]
        UC04[Đăng nhập]
    end

    subgraph CustomerArea[Khu vực khách hàng]
        UC05[Chọn suất chiếu]
        UC06[Chọn ghế]
        UC07[Xác nhận đặt vé]
        UC08[Xem vé đã đặt]
    end

    subgraph AdminArea[Khu vực quản trị]
        UC09[Quản lý phim]
        UC10[Quản lý phòng chiếu]
        UC11[Quản lý lịch chiếu]
        UC12[Lọc lịch chiếu]
    end

    Customer --> UC01
    Customer --> UC02
    Customer --> UC03
    Customer --> UC04
    Customer --> UC05
    Customer --> UC06
    Customer --> UC07
    Customer --> UC08
    Admin --> UC04
    Admin --> UC09
    Admin --> UC10
    Admin --> UC11
    Admin --> UC12
""")

    doc.add_heading("2.3 Sơ đồ luồng tổng thể", level=2)
    add_mermaid(doc, "2.3.1 Mermaid - Luồng nghiệp vụ tổng thể", """
flowchart TD
    A[Người dùng truy cập trang chủ] --> B{Đã đăng nhập?}
    B -- Chưa --> C[Xem phim công khai]
    C --> D[Đăng ký hoặc đăng nhập]
    B -- Rồi --> E{Vai trò}
    D --> E
    E -- ROLE_CUSTOMER --> F[Xem chi tiết phim]
    F --> G[Chọn suất chiếu SELLING]
    G --> H[Chọn ghế còn trống]
    H --> I{Ghế hợp lệ?}
    I -- Không --> H
    I -- Có --> J[Tạo Booking]
    J --> K[Hiển thị xác nhận]
    K --> L[Xem lịch sử đặt vé]
    E -- ROLE_ADMIN --> M[Quản lý phim/phòng/lịch chiếu]
    M --> N[Lưu dữ liệu]
    N --> O[Lọc và xem lịch chiếu]
""")

    doc.add_heading("2.4 Sơ đồ chuyển trạng thái", level=2)
    doc.add_paragraph("Trạng thái suất chiếu hiện gồm: SCHEDULED, SELLING, SOLD_OUT và CANCELLED.")
    add_mermaid(doc, "2.4.1 Mermaid - Trạng thái suất chiếu", """
stateDiagram-v2
    [*] --> SCHEDULED : Tạo lịch chiếu
    SCHEDULED --> SELLING : Mở bán vé
    SELLING --> SOLD_OUT : Hết ghế
    SCHEDULED --> CANCELLED : Hủy lịch
    SELLING --> CANCELLED : Hủy suất chiếu
    SOLD_OUT --> CANCELLED : Hủy sau khi đã bán hết
    CANCELLED --> [*]
""")

    doc.add_heading("2.5 Phân quyền", level=2)
    add_table(
        doc,
        ["STT", "Chức năng", "Chưa đăng nhập", "Khách hàng", "Quản trị viên"],
        [
            ("1", "Trang chủ, danh sách phim", "Được phép", "Được phép", "Tự chuyển /showtimes"),
            ("2", "Đăng ký", "Được phép", "Không cần thiết", "Không cần thiết"),
            ("3", "Đăng nhập", "Được phép", "Không cần thiết", "Không cần thiết"),
            ("4", "Đặt vé", "Không được phép", "Được phép", "Không áp dụng"),
            ("5", "Vé đã đặt", "Không được phép", "Được phép", "Không áp dụng"),
            ("6", "Quản lý phim", "Không được phép", "Không được phép", "Được phép"),
            ("7", "Quản lý phòng chiếu", "Không được phép", "Không được phép", "Được phép"),
            ("8", "Quản lý lịch chiếu", "Không được phép", "Không được phép", "Được phép"),
        ],
        [Cm(1.2), Cm(4.2), Cm(3.6), Cm(3.5), Cm(4)],
    )
    doc.add_paragraph(
        "Phiên bản hiện tại sử dụng Spring Security để kiểm soát truy cập theo URL. "
        "Các URL /movies/**, /auditoriums/** và /showtimes/** yêu cầu ROLE_ADMIN; "
        "/booking/** và /my-bookings yêu cầu ROLE_CUSTOMER."
    )

    doc.add_heading("2.6 Site Map", level=2)
    add_mermaid(doc, "2.6.1 Mermaid - Site Map", """
flowchart TD
    Home["/ - Trang chủ"]
    Login["/login"]
    Register["/register"]
    MovieDetail["/movie/:id"]
    BookingPage["/booking/:showtimeId"]
    Confirm["/booking/:showtimeId/confirm"]
    MyBookings["/my-bookings"]
    AdminShowtimes["/showtimes"]
    AdminMovies["/movies"]
    AdminAuditoriums["/auditoriums"]

    Home --> MovieDetail
    Home --> Login
    Home --> Register
    MovieDetail --> BookingPage
    BookingPage --> Confirm
    Confirm --> MyBookings
    Login --> Home
    Login --> AdminShowtimes
    AdminShowtimes --> AdminMovies
    AdminShowtimes --> AdminAuditoriums
    AdminShowtimes --> AdminShowtimes
""")

    doc.add_heading("PHẦN 3: CHỨC NĂNG", level=1)
    add_table(
        doc,
        ["STT", "Mã", "Tên chức năng", "Phân hệ", "Tác nhân"],
        [
            ("1", "UC-01", "Đăng ký tài khoản khách hàng", "Tài khoản", "Khách hàng"),
            ("2", "UC-02", "Đăng nhập", "Tài khoản", "Khách hàng, quản trị viên"),
            ("3", "UC-03", "Xem phim và chi tiết phim", "Khách hàng", "Khách hàng"),
            ("4", "UC-04", "Đặt vé và chọn ghế", "Đặt vé", "Khách hàng"),
            ("5", "UC-05", "Xem vé đã đặt", "Đặt vé", "Khách hàng"),
            ("6", "UC-06", "Quản lý phim", "Quản trị", "Quản trị viên"),
            ("7", "UC-07", "Quản lý phòng chiếu", "Quản trị", "Quản trị viên"),
            ("8", "UC-08", "Quản lý và lọc lịch chiếu", "Quản trị", "Quản trị viên"),
        ],
        [Cm(1.2), Cm(1.7), Cm(5), Cm(4), Cm(4.6)],
    )

    use_cases = [
        (1, "UC-01", "Đăng ký tài khoản khách hàng", "Khách hàng",
         "Người dùng chọn Đăng ký trên thanh điều hướng.",
         "Người dùng chưa đăng nhập và chưa có tài khoản trùng username/email.",
         "Tài khoản mới được lưu với vai trò ROLE_CUSTOMER và mật khẩu đã mã hóa.",
         ["Người dùng mở trang /register.",
          "Hệ thống hiển thị form username, họ tên, email, mật khẩu và xác nhận mật khẩu.",
          "Người dùng nhập thông tin và gửi form.",
          "Hệ thống kiểm tra độ dài mật khẩu, mật khẩu xác nhận, username và email trùng.",
          "Nếu hợp lệ, hệ thống mã hóa mật khẩu bằng PasswordEncoder và lưu User.",
          "Hệ thống hiển thị thông báo thành công và chuyển về trang đăng nhập."],
         [("username", "Tên đăng nhập", "Text", "Có", "Duy nhất trong hệ thống."),
          ("fullName", "Họ tên", "Text", "Có", "Tên hiển thị của khách hàng."),
          ("email", "Email", "Email", "Có", "Duy nhất trong hệ thống."),
          ("password", "Mật khẩu", "Password", "Có", "Tối thiểu 6 ký tự, lưu dạng hash."),
          ("confirmPassword", "Nhập lại mật khẩu", "Password", "Có", "Phải khớp với mật khẩu.")]),
        (2, "UC-02", "Đăng nhập", "Khách hàng, quản trị viên",
         "Người dùng truy cập /login hoặc một trang cần xác thực.",
         "Tài khoản đã tồn tại trong hệ thống.",
         "Người dùng được xác thực và điều hướng theo vai trò.",
         ["Người dùng mở trang /login.",
          "Hệ thống hiển thị form đăng nhập.",
          "Người dùng nhập username và password.",
          "Spring Security xác thực bằng UserDetailsService.",
          "ROLE_ADMIN được chuyển đến /showtimes; ROLE_CUSTOMER được chuyển đến /.",
          "Nếu sai thông tin, hệ thống quay lại /login?error=true và hiển thị lỗi."],
         [("username", "Tên đăng nhập", "Text", "Có", "Tên đăng nhập đã đăng ký."),
          ("password", "Mật khẩu", "Password", "Có", "Mật khẩu tương ứng với tài khoản.")]),
        (3, "UC-03", "Xem phim và chi tiết phim", "Khách hàng",
         "Người dùng truy cập trang chủ hoặc chọn một phim.",
         "Hệ thống có dữ liệu phim.",
         "Người dùng xem được thông tin phim và các suất chiếu đang SELLING.",
         ["Người dùng mở trang /.",
          "Hệ thống lấy danh sách phim sắp xếp theo tên.",
          "Người dùng chọn phim để xem chi tiết.",
          "Hệ thống tải Movie theo id.",
          "Hệ thống hiển thị thông tin phim và danh sách suất chiếu có trạng thái SELLING.",
          "Người dùng chọn suất chiếu để đặt vé nếu đã đăng nhập khách hàng."],
         [("id", "Mã phim", "Path", "Có", "Dùng khi mở /movie/{id}."),
          ("title", "Tên phim", "Text", "Không nhập", "Hiển thị từ dữ liệu Movie."),
          ("showtimes", "Suất chiếu", "List", "Không nhập", "Chỉ hiển thị suất chiếu đang bán.")]),
        (4, "UC-04", "Đặt vé và chọn ghế", "Khách hàng",
         "Khách hàng chọn một suất chiếu.",
         "Khách hàng đã đăng nhập; suất chiếu tồn tại.",
         "Booking được tạo với danh sách ghế, tổng tiền và ngày đặt.",
         ["Khách hàng mở /booking/{showtimeId}.",
          "Hệ thống lấy thông tin suất chiếu và danh sách ghế đã được đặt.",
          "Khách hàng chọn một hoặc nhiều ghế còn trống.",
          "Hệ thống kiểm tra phải chọn ít nhất một ghế.",
          "Hệ thống kiểm tra ghế chưa bị đặt bởi booking khác cùng suất chiếu.",
          "Hệ thống tính tổng tiền bằng ticketPrice nhân số ghế và lưu Booking.",
          "Hệ thống hiển thị trang xác nhận đặt vé."],
         [("showtimeId", "Mã suất chiếu", "Path", "Có", "Suất chiếu được đặt vé."),
          ("seats", "Danh sách ghế", "Text/List", "Có", "Danh sách ghế dạng A1,A2."),
          ("totalPrice", "Tổng tiền", "Number", "Tự tính", "ticketPrice * số ghế."),
          ("bookingDate", "Ngày đặt", "Datetime", "Tự sinh", "Thời điểm tạo booking.")]),
        (5, "UC-05", "Xem vé đã đặt", "Khách hàng",
         "Khách hàng chọn mục Vé của tôi hoặc truy cập /my-bookings.",
         "Khách hàng đã đăng nhập.",
         "Danh sách booking của khách hàng được hiển thị theo ngày đặt giảm dần.",
         ["Khách hàng mở /my-bookings.",
          "Hệ thống lấy User hiện tại theo username.",
          "Hệ thống truy vấn Booking theo User và sắp xếp bookingDate giảm dần.",
          "Hệ thống hiển thị thông tin phim, suất chiếu, ghế và tổng tiền.",
          "Nếu chưa có booking, giao diện hiển thị trạng thái rỗng phù hợp."],
         [("user", "Khách hàng", "Object", "Tự xác định", "Lấy từ tài khoản đang đăng nhập."),
          ("bookings", "Danh sách vé", "List", "Không nhập", "Dữ liệu trả về từ BookingService.")]),
        (6, "UC-06", "Quản lý phim", "Quản trị viên",
         "Quản trị viên mở /movies.",
         "Quản trị viên đã đăng nhập với ROLE_ADMIN.",
         "Danh sách phim được thêm, cập nhật hoặc xóa.",
         ["Quản trị viên mở trang quản lý phim.",
          "Hệ thống hiển thị danh sách phim và form nhập phim.",
          "Quản trị viên nhập tên phim, thể loại, thời lượng, phân loại tuổi, ngày phát hành, poster và mô tả.",
          "Nếu có file poster, hệ thống lưu vào uploads/movies với tên UUID.",
          "Hệ thống lưu Movie và tải lại danh sách.",
          "Quản trị viên có thể chọn sửa hoặc xóa phim trong danh sách."],
         [("title", "Tên phim", "Text", "Có", "Tên phim hiển thị trên trang chủ."),
          ("genre", "Thể loại", "Text", "Có", "Ví dụ: hành động, hoạt hình, tâm lý."),
          ("durationMinutes", "Thời lượng", "Number", "Có", "Số phút chiếu phim."),
          ("ageRating", "Phân loại tuổi", "Text", "Không", "Ví dụ: P, T13, T16, T18."),
          ("releaseDate", "Ngày phát hành", "Date", "Không", "Ngày phim phát hành."),
          ("posterFile/posterUrl", "Poster", "File/Text", "Không", "Ảnh poster hoặc URL đã lưu."),
          ("description", "Mô tả", "Textarea", "Không", "Tối đa theo entity hiện tại 2000 ký tự.")]),
        (7, "UC-07", "Quản lý phòng chiếu", "Quản trị viên",
         "Quản trị viên mở /auditoriums.",
         "Quản trị viên đã đăng nhập với ROLE_ADMIN.",
         "Danh sách phòng chiếu được cập nhật.",
         ["Quản trị viên mở trang quản lý phòng chiếu.",
          "Hệ thống hiển thị danh sách phòng và form nhập.",
          "Quản trị viên nhập tên phòng, sức chứa và định dạng màn hình.",
          "Hệ thống kiểm tra tên phòng không trùng khi thêm hoặc sửa.",
          "Nếu hợp lệ, hệ thống lưu Auditorium.",
          "Quản trị viên có thể sửa hoặc xóa phòng chiếu."],
         [("name", "Tên phòng", "Text", "Có", "Duy nhất, ví dụ: Phòng A."),
          ("capacity", "Sức chứa", "Number", "Có", "Số ghế của phòng chiếu."),
          ("screenType", "Định dạng màn hình", "Text", "Có", "Ví dụ: 2D, 3D, IMAX.")]),
        (8, "UC-08", "Quản lý và lọc lịch chiếu", "Quản trị viên",
         "Quản trị viên mở /showtimes hoặc thao tác thêm/sửa/xóa/lọc.",
         "Quản trị viên đã đăng nhập; hệ thống có dữ liệu phim và phòng chiếu.",
         "Lịch chiếu được tạo, cập nhật, xóa hoặc hiển thị theo điều kiện lọc.",
         ["Quản trị viên mở trang /showtimes.",
          "Hệ thống hiển thị bộ lọc ngày, phim, phòng và danh sách lịch chiếu.",
          "Khi thêm/sửa, quản trị viên chọn phim, phòng, thời gian bắt đầu, giá vé, trạng thái và ghi chú.",
          "Hệ thống kiểm tra giá vé lớn hơn 0 và tham chiếu phim/phòng hợp lệ.",
          "Nếu hợp lệ, hệ thống lưu Showtime.",
          "Khi lọc, hệ thống truy vấn theo khoảng thời gian trong ngày và các điều kiện phim/phòng tùy chọn."],
         [("movieId", "Phim", "Select", "Có", "Tham chiếu Movie hợp lệ."),
          ("auditoriumId", "Phòng chiếu", "Select", "Có", "Tham chiếu Auditorium hợp lệ."),
          ("startsAt", "Thời gian bắt đầu", "Datetime", "Có", "Ngày giờ bắt đầu suất chiếu."),
          ("ticketPrice", "Giá vé", "Currency", "Có", "Phải lớn hơn 0."),
          ("status", "Trạng thái", "Select", "Có", "SCHEDULED, SELLING, SOLD_OUT, CANCELLED."),
          ("note", "Ghi chú", "Textarea", "Không", "Thông tin bổ sung cho suất chiếu."),
          ("date/movieId/auditoriumId", "Bộ lọc", "Query", "Không", "Lọc danh sách lịch chiếu.")]),
    ]
    for uc in use_cases:
        add_use_case(doc, *uc)

    doc.add_heading("PHẦN 4: CÁC COMPONENT, THÔNG BÁO, CẢNH BÁO", level=1)
    doc.add_heading("4.1 Component giao diện", level=2)
    add_table(
        doc,
        ["STT", "Component", "Mô tả", "Màn hình sử dụng"],
        [
            ("1", "Topbar/Menu", "Điều hướng giữa trang chủ, phim, vé của tôi, quản trị, đăng nhập/đăng ký/đăng xuất.", "Toàn hệ thống"),
            ("2", "Auth Form", "Form đăng nhập và đăng ký tài khoản.", "/login, /register"),
            ("3", "Movie List/Card", "Hiển thị danh sách phim và poster.", "/"),
            ("4", "Movie Detail", "Hiển thị mô tả phim và danh sách suất chiếu đang bán.", "/movie/{id}"),
            ("5", "Seat Picker", "Cho phép khách hàng chọn ghế và tránh ghế đã được đặt.", "/booking/{showtimeId}"),
            ("6", "Booking Confirmation", "Hiển thị kết quả đặt vé thành công.", "/booking/{showtimeId}/confirm"),
            ("7", "Filter Bar", "Bộ lọc lịch chiếu theo ngày, phim và phòng chiếu.", "/showtimes"),
            ("8", "Data Table", "Bảng hiển thị danh sách phim, phòng chiếu, lịch chiếu và booking.", "Các màn hình danh sách"),
            ("9", "Alert Message", "Thông báo thao tác thành công hoặc lỗi validate.", "Các màn hình form"),
        ],
        [Cm(1.2), Cm(3.4), Cm(7), Cm(4.9)],
    )

    doc.add_heading("4.2 Thông báo hệ thống", level=2)
    add_table(
        doc,
        ["STT", "Tình huống", "Thông báo"],
        [
            ("1", "Đăng ký thành công", "Đăng ký tài khoản thành công! Vui lòng đăng nhập."),
            ("2", "Đăng nhập thất bại", "Tên đăng nhập hoặc mật khẩu không chính xác."),
            ("3", "Chưa chọn ghế", "Vui lòng chọn ít nhất một ghế."),
            ("4", "Ghế đã có người đặt", "Ghế {seat} đã có người đặt trước."),
            ("5", "Lưu phim thành công", "Đã lưu phim thành công."),
            ("6", "Cập nhật phim thành công", "Đã cập nhật phim thành công."),
            ("7", "Xóa phim thành công", "Đã xóa phim thành công."),
            ("8", "Lưu phòng chiếu thành công", "Đã lưu phòng chiếu thành công."),
            ("9", "Tên phòng trùng", "Phòng chiếu '{name}' đã tồn tại."),
            ("10", "Lưu lịch chiếu thành công", "Đã lưu lịch chiếu thành công."),
            ("11", "Giá vé không hợp lệ", "Giá vé phải lớn hơn 0."),
        ],
        [Cm(1.2), Cm(5.4), Cm(9.9)],
    )

    doc.add_heading("4.3 Cảnh báo và validate", level=2)
    validations = [
        "Username và email phải duy nhất khi đăng ký tài khoản.",
        "Mật khẩu phải có tối thiểu 6 ký tự và xác nhận mật khẩu phải khớp.",
        "Khách hàng phải chọn ít nhất một ghế khi đặt vé.",
        "Ghế được chọn không được trùng với ghế đã có booking trong cùng suất chiếu.",
        "Tên phòng chiếu không được trùng khi thêm mới hoặc cập nhật.",
        "Giá vé của suất chiếu phải lớn hơn 0.",
        "Các chức năng quản trị chỉ cho phép ROLE_ADMIN truy cập.",
        "Các chức năng đặt vé và xem vé đã đặt chỉ cho phép ROLE_CUSTOMER truy cập.",
    ]
    for item in validations:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("PHẦN 5: LINK ISSUE", level=1)
    doc.add_paragraph("Tạm thời hệ thống chưa liên kết công cụ quản lý issue/Jira.")
    add_table(
        doc,
        ["STT", "Mã issue", "Nội dung", "Trạng thái"],
        [("1", "N/A", "Chưa có issue liên kết cho giai đoạn hiện tại.", "Tạm thời")],
        [Cm(1.2), Cm(3), Cm(8), Cm(4.3)],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("SRS-CGV-CINEMA-001 | Phiên bản 1.0")

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
